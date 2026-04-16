"""
Reads chapter2_full.md, parses it into structured blocks, and appends the
content to thesis/chapter1/chapter1_full.docx using python-docx, preserving
headings, bullet lists, bold text, code blocks, and horizontal rules.
"""
import os
import re
import copy
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Safe style lookup ─────────────────────────────────────────────────────────
def safe_style(doc, *names):
    """Return the first style from `names` that exists, else 'Normal'."""
    for name in names:
        try:
            return doc.styles[name].name
        except KeyError:
            pass
    return 'Normal'

# ── Paths ─────────────────────────────────────────────────────────────────────
BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
MD_PATH   = os.path.join(BASE, "thesis", "chapter3_full.md")
DOCX_PATH = os.path.join(BASE, "thesis", "chapter1", "chapter1_full.docx")

# ── Load files ─────────────────────────────────────────────────────────────────
doc = Document(DOCX_PATH)

with open(MD_PATH, encoding="utf-8") as f:
    md_text = f.read()

# ── Helper: add a horizontal rule paragraph ────────────────────────────────────
def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'AAAAAA')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ── Helper: add inline-styled paragraph (bold markers **text**) ────────────────
def add_styled_paragraph(doc, text, style_name='Normal'):
    p = doc.add_paragraph(style=style_name)
    # Split on **...** markers
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            run = p.add_run(part[1:-1])
            run.italic = True
        else:
            p.add_run(part)
    return p

# ── Helper: add a code block ───────────────────────────────────────────────────
def add_code_block(doc, code_text):
    for line in code_text.split('\n'):
        p = doc.add_paragraph(style='Normal')
        # Remove paragraph spacing
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after  = Pt(0)
        run = p.add_run(line if line else ' ')
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        # Light grey background via paragraph shading
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F0F0F0')
        pPr.append(shd)

# ── Page break before Chapter 2 ────────────────────────────────────────────────
doc.add_page_break()

# ── Parse and render markdown ──────────────────────────────────────────────────
lines = md_text.split('\n')
i = 0
while i < len(lines):
    line = lines[i]

    # ── Code block ──────────────────────────────────────────────────────────
    if line.strip().startswith('```'):
        i += 1
        code_lines = []
        while i < len(lines) and not lines[i].strip().startswith('```'):
            code_lines.append(lines[i])
            i += 1
        add_code_block(doc, '\n'.join(code_lines))
        i += 1  # skip closing ```
        continue

    # ── Horizontal rule ─────────────────────────────────────────────────────
    if line.strip() == '---':
        add_horizontal_rule(doc)
        i += 1
        continue

    # ── Headings ────────────────────────────────────────────────────────────
    heading_match = re.match(r'^(#{1,4})\s+(.*)', line)
    if heading_match:
        level = len(heading_match.group(1))
        text  = heading_match.group(2).strip()
        # Map markdown heading levels to Word heading levels
        # # → Heading 1, ## → Heading 2, ### → Heading 3
        word_level = min(level, 3)
        doc.add_heading(text, level=word_level)
        i += 1
        continue

    # ── Bullet list item ─────────────────────────────────────────────────────
    bullet_match = re.match(r'^(\s*)[*\-]\s+(.*)', line)
    numbered_match = re.match(r'^(\s*)\d+\.\s+(.*)', line)
    if bullet_match:
        text = bullet_match.group(2)
        p = doc.add_paragraph(style=safe_style(doc, 'List Bullet', 'Normal'))
        parts = re.split(r'(\*\*[^*]+\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                p.add_run(part)
        i += 1
        continue
    if numbered_match:
        text = numbered_match.group(2)
        p = doc.add_paragraph(style=safe_style(doc, 'List Number', 'Normal'))
        parts = re.split(r'(\*\*[^*]+\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                p.add_run(part)
        i += 1
        continue

    # ── Inline code only line (skip mermaid blocks) ─────────────────────────

    # ── Empty line → skip ────────────────────────────────────────────────────
    if not line.strip():
        i += 1
        continue

    # ── Regular paragraph ────────────────────────────────────────────────────
    # Strip leading > (blockquote) if any
    text = re.sub(r'^>\s?', '', line)
    # Remove inline backtick code markers for cleaner rendering
    text = re.sub(r'`([^`]+)`', r'\1', text)
    # Remove italic/bold markers for clean text (bold handled below)
    add_styled_paragraph(doc, text, style_name='Normal')
    i += 1

# ── Save ────────────────────────────────────────────────────────────────────────
doc.save(DOCX_PATH)
print(f"Done! Chapter 2 appended to:\n   {DOCX_PATH}")
